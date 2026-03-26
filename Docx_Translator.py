# docx_Translator
import os
import re
import unicodedata
import json
import sys
import threading
import queue
import deepl # to translate text with contextual accuracy
import time
import traceback
import webbrowser
import asyncio
import tkinter as tk

from googletrans import Translator # to translate text
from pathlib import Path
from tkinter import Tk
from tkinter import filedialog as fd
from tkinter import messagebox
from tkinter import ttk
from dotenv import load_dotenv # to load environment variables from .env file
from zipfile import BadZipFile # to handle invalid .docx files
from docx import Document   # to read and write .docx files
from docx.table import Table
from docx.text.paragraph import Paragraph
from datetime import datetime


# load the .env file from the same directory as this script
try:
    env_path = Path(__file__).parent / "Project_env.env"
except NameError:
    env_path = Path.cwd() / "Project_env.env"

load_dotenv(env_path) # load environment variables from .env file

# define rate limit parameters
rate_limit_minute = 450  # translator plan rate limit: 450 requests per minute
delay_between_requests = 60/rate_limit_minute  # calculate delay between requests in seconds
max_retries = 5 # maximum number of retries for failed requests
try:
    _base_dir = Path(__file__).parent
except NameError:
    _base_dir = Path.cwd()

ERROR_LOG_DIR = os.getenv("lin_error_logs_dir") or os.getenv("error_logs_dir") or str(_base_dir)

# create google translator object
TRANSLATION_PROVIDER = "deepl"  # change to "deepl" to use DeepL
TRANSLATION_MEMORY_ENABLED = True
TRANSLATION_MEMORY_PATH = None
translation_memory = {}

translator = Translator()
deepl_translator = None

def set_translator(new_translator):
    global translator
    translator = new_translator

async def close_translator_client(translator_obj):
    try:
        if hasattr(translator_obj, "client"):
            await translator_obj.client.aclose()
    except Exception:
        pass

def get_deepl_translator():
    global deepl_translator
    if deepl_translator is not None:
        return deepl_translator

    auth_key = os.getenv("deepL_auth_key")
    if not auth_key:
        raise RuntimeError("DeepL API key not found. Set DEEPL_API_KEY in Project_env.env.")

    deepl_translator = deepl.Translator(auth_key)
    return deepl_translator

async def provider_translate(text, target_lang):
    if TRANSLATION_PROVIDER == "deepl":
        if TRANSLATION_MEMORY_ENABLED:
            key = f"{target_lang}||{text}"
            cached = translation_memory.get(key)
            if cached is not None:
                return cached

        translator_obj = get_deepl_translator()
        result = await asyncio.to_thread(
            translator_obj.translate_text,
            text,
            target_lang=target_lang
        )
        translated = result.text or ""
        if TRANSLATION_MEMORY_ENABLED:
            translation_memory[key] = translated
        return translated

    # default to googletrans
    result = await translator.translate(text, dest=target_lang)
    return result.text

NON_TRANSLATABLE_PATTERNS = [
    re.compile(r"https?://[^\s]+|www\.[^\s]+", re.IGNORECASE),  # URLs
    re.compile(r"\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b"),  # emails
    re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"),  # dates
    re.compile(r"[$€£¥]\s?\d{1,3}(?:,\d{3})*(?:\.\d+)?|\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s?[$€£¥]\b"),  # currency
    re.compile(r"\b\d+(?:[.,]\d+)?%"),  # percentages
    re.compile(r"\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b"),  # numbers
]

PROPER_NOUN_SEQUENCE_RE = re.compile(
    r"\b(?:[A-Z][a-z]+(?:['-][A-Z][a-z]+)?)(?:\s+(?:[A-Z][a-z]+(?:['-][A-Z][a-z]+)?))*\b"
)
ALL_CAPS_RE = re.compile(r"\b[A-Z]{2,}\b")
PLACEHOLDER_RE = re.compile(r"__NTX_\d+__")

COMMON_ENGLISH_WORDS = {
    "the","and","of","to","in","for","with","on","at","from","by","is","are","was",
    "were","be","this","that","it","as","an","a","or","but","not","we","you","they",
    "he","she","them","his","her","their","our","us"
}

def load_known_translations():
    try:
        rules_path = Path(__file__).parent / "Files/quality_rules.json"
    except NameError:
        rules_path = Path.cwd() / "Files/quality_rules.json"

    try:
        with open(rules_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        translations = data.get("known_translations", {})
        if isinstance(translations, dict):
            return translations
    except Exception:
        pass

    return {
        "hello": "hola",
        "good morning": "buenos días",
        "good afternoon": "buenas tardes",
        "good night": "buenas noches"
    }

KNOWN_TRANSLATIONS = load_known_translations()

def load_protected_words():
    try:
        words_path = Path(__file__).parent / "Files/protected_words.json"
    except NameError:
        words_path = Path.cwd() / "Files/protected_words.json"

    try:
        with open(words_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        words = data.get("protected_words", [])
        if isinstance(words, list):
            return [str(w) for w in words if str(w).strip()]
    except Exception:
        pass

    return [
        "Ana",
        "Maria",
        "María",
        "Jose",
        "José",
        "Juan",
        "Luis",
        "Marta",
        "Carlos",
        "Pedro",
        "Lucia",
        "Lucía",
        "Sofia",
        "Sofía",
        "Andrea",
        "Diego",
        "Miguel",
        "Pablo",
    ]

PROTECTED_WORDS = load_protected_words()

# define a variable to select the file to be translated
def select_docx_file():
    """Open a file dialog and return the selected filepath to translate"""
    # create hidden root window
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    # create child window(file dialog)
    file_path = fd.askopenfilename( # assign a variable to open a dialog and select the file
        title="Choose a file to translate",
        filetypes=[
            ("Word Documents", "*.docx"), # shows only .docx files
            ("All FIles", "*.*") # show every type of file
        ],
        # use the environment variable to set the initial directory and a spare default value
        initialdir=os.getenv("lin_test_docs_dir", os.getenv("app_tools_dir"))
    )

    # destroy the root dialog window
    root.destroy()

    # Return None instead of empty string for better logic
    return file_path if file_path else None

def load_settings():
    try:
        settings_path = Path(__file__).parent / "settings.json"
    except NameError:
        settings_path = Path.cwd() / "settings.json"

    defaults = {
        "api_key": "",
        "source_lang": "EN",
        "target_lang": "ES",
        "open_settings_on_start": True,
        "use_gui": True,
        "ui_language": "EN",
        "translation_memory_file": "translation_memory.json"
    }

    if not settings_path.exists():
        try:
            with open(settings_path, "w", encoding="utf-8") as f:
                json.dump(defaults, f, indent=2, ensure_ascii=False)
        except Exception:
            pass
        return defaults, settings_path

    try:
        with open(settings_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            merged = defaults.copy()
            merged.update(data)
            return merged, settings_path
    except Exception:
        pass

    return defaults, settings_path

def init_translation_memory(settings, base_dir):
    global TRANSLATION_MEMORY_PATH, translation_memory
    filename = settings.get("translation_memory_file") or "translation_memory.json"
    TRANSLATION_MEMORY_PATH = str(base_dir / filename)

    if TRANSLATION_PROVIDER != "deepl":
        translation_memory = {}
        return

    try:
        if os.path.exists(TRANSLATION_MEMORY_PATH):
            with open(TRANSLATION_MEMORY_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                translation_memory = data
    except Exception:
        translation_memory = {}

def save_translation_memory():
    if TRANSLATION_PROVIDER != "deepl":
        return
    if not TRANSLATION_MEMORY_PATH:
        return
    try:
        with open(TRANSLATION_MEMORY_PATH, "w", encoding="utf-8") as f:
            json.dump(translation_memory, f, indent=2, ensure_ascii=False)
    except Exception:
        pass

def prompt_edit_settings(settings_path):
    try:
        answer = input(f"Edit settings now? ({settings_path}) (y/n): ").strip().lower()
        if answer == "y":
            if sys.platform.startswith("win"):
                os.system(f'start "" "{settings_path}"')
            elif sys.platform == "darwin":
                os.system(f'open "{settings_path}"')
            else:
                os.system(f'xdg-open "{settings_path}"')
    except Exception:
        pass

def select_output_folder():
    try:
        return fd.askdirectory(title="Choose output folder")
    except Exception:
        return ""

def open_settings_file(settings_path):
    try:
        if sys.platform.startswith("win"):
            os.system(f'start "" "{settings_path}"')
        elif sys.platform == "darwin":
            os.system(f'open "{settings_path}"')
        else:
            os.system(f'xdg-open "{settings_path}"')
    except Exception:
        pass

def save_settings(settings, settings_path):
    try:
        with open(settings_path, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2, ensure_ascii=False)
    except Exception:
        pass

UI_STRINGS = {
    "EN": {
        "app_title": "Docx Translator",
        "language_prompt": "Choose interface language",
        "english": "English",
        "spanish": "Spanish",
        "select_files": "Select .docx files to translate...",
        "add_files": "Add Files",
        "clear_list": "Clear List",
        "settings": "Settings",
        "start": "Start",
        "stop": "Stop",
        "summary": "Summary",
        "no_files": "Please add at least one .docx file.",
        "batch_start": "Batch start: {count} file(s).",
        "processing_file": "Processing file {idx}/{total}: {file}",
        "batch_completed": "Batch completed.",
        "all_files_ok": "All files translated successfully.",
        "open_folder": "Open translated files folder?",
        "summary_output": "Output",
        "summary_paragraphs": "Paragraphs",
        "summary_tables": "Tables",
        "summary_images": "Images",
        "summary_headers": "Headers/Footers",
        "summary_time": "Time",
        "quality_warnings": "Quality warnings",
        "warn_length_ratio": "Paragraph #{idx} length ratio suspicious (src {src} vs out {out}).",
        "warn_english_words": "Paragraph #{idx} contains common English words.",
        "warn_known_translation": "Paragraph #{idx} may have missed known translation for '{phrase}'.",
        "settings_title": "Settings",
        "target_language": "Target Language",
        "open_settings_start": "Open settings on start",
        "use_gui": "Use GUI (disable to use CLI)",
        "save": "Save",
        "cancel": "Cancel",
        "file_idle": "File: idle",
        "ready": "Ready."
    },
    "ES": {
        "app_title": "Traductor Docx",
        "language_prompt": "Elige el idioma de la interfaz",
        "english": "Inglés",
        "spanish": "Español",
        "select_files": "Seleccione archivos .docx para traducir...",
        "add_files": "Agregar archivos",
        "clear_list": "Limpiar lista",
        "settings": "Configuración",
        "start": "Iniciar",
        "stop": "Detener",
        "summary": "Resumen",
        "no_files": "Agregue al menos un archivo .docx.",
        "batch_start": "Inicio de lote: {count} archivo(s).",
        "processing_file": "Procesando archivo {idx}/{total}: {file}",
        "batch_completed": "Lote completado.",
        "all_files_ok": "Todos los archivos fueron traducidos correctamente.",
        "open_folder": "¿Abrir la carpeta de archivos traducidos?",
        "summary_output": "Salida",
        "summary_paragraphs": "Párrafos",
        "summary_tables": "Tablas",
        "summary_images": "Imágenes",
        "summary_headers": "Encabezados/Pies",
        "summary_time": "Tiempo",
        "quality_warnings": "Advertencias de calidad",
        "warn_length_ratio": "El párrafo #{idx} tiene una proporción sospechosa (orig {src} vs trad {out}).",
        "warn_english_words": "El párrafo #{idx} contiene palabras comunes en inglés.",
        "warn_known_translation": "El párrafo #{idx} pudo omitir la traducción conocida de '{phrase}'.",
        "settings_title": "Configuración",
        "target_language": "Idioma de salida",
        "open_settings_start": "Abrir configuración al inicio",
        "use_gui": "Usar GUI (deshabilitar para CLI)",
        "save": "Guardar",
        "cancel": "Cancelar",
        "file_idle": "Archivo: en espera",
        "ready": "Listo."
    }
}

def get_ui_strings(settings):
    lang = (settings.get("ui_language") or "EN").upper()
    return UI_STRINGS.get(lang, UI_STRINGS["EN"])

def show_language_selector(settings, settings_path):
    lang_root = Tk()
    lang_root.title("Language")
    lang_root.geometry("360x180")
    lang_root.resizable(False, False)

    label = ttk.Label(lang_root, text=UI_STRINGS["EN"]["language_prompt"])
    label.pack(pady=15)

    choice = {"lang": settings.get("ui_language", "EN")}

    def set_lang(lang_code):
        choice["lang"] = lang_code
        settings["ui_language"] = lang_code
        save_settings(settings, settings_path)
        lang_root.destroy()

    btn_frame = ttk.Frame(lang_root)
    btn_frame.pack(pady=10)

    ttk.Button(btn_frame, text=UI_STRINGS["EN"]["english"], command=lambda: set_lang("EN")).pack(side="left", padx=10)
    ttk.Button(btn_frame, text=UI_STRINGS["EN"]["spanish"], command=lambda: set_lang("ES")).pack(side="left", padx=10)

    lang_root.mainloop()

def validate_language_code(value, fallback="ES"):
    if not isinstance(value, str):
        return fallback
    code = value.strip().upper()
    if re.fullmatch(r"[A-Z]{2,3}", code):
        return code
    print(f"Warning: Invalid language code '{value}'. Using '{fallback}'.")
    log_error("Invalid language code", value)
    return fallback

def prompt_add_another_file():
    try:
        answer = input("Add another file? (y/n): ").strip().lower()
        return answer == "y"
    except Exception:
        return False

def ensure_error_log_dir():
    try:
        os.makedirs(ERROR_LOG_DIR, exist_ok=True)
    except Exception:
        pass

def log_error(context, error, file_path=None):
    ensure_error_log_dir()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_name = "error_log.txt"
    log_path = os.path.join(ERROR_LOG_DIR, log_name)

    details = [
        f"[{timestamp}] {context}",
        f"Error: {repr(error)}",
    ]
    if file_path:
        details.append(f"File: {file_path}")
    details.append("Traceback:")
    details.append(traceback.format_exc())
    details.append("-" * 60)

    try:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write("\n".join(details) + "\n")
    except Exception:
        pass

def file_validation(file_path):
    """Validate if a file path was selected"""
    if file_path is None: # when no file is selected
        return
    elif not file_path.lower().endswith(".docx"): # validate file extension
        print("\nError: The selected file must be a '.docx' file.")
        print("Tip: Save or export your document as .docx and try again.")
        return
    else:
        try: # validate file existence
            # When file is selected
            print(f"\nFile selected to translate: {file_path}",
                    f"\nFile path: {os.path.dirname(file_path)}",
                    f"\nFile name: {os.path.basename(file_path)}",
                    f"\nFile size: {os.path.getsize(file_path)} KB", sep="")
            return True
        except FileExistsError: # file does not exist
            print(f"\nError: The file {file_path} does not exist.")
            print("Tip: Verify the file path and try again.")
            log_error("File not found", "FileExistsError", file_path)
            return
        except BadZipFile as e: # file is not a valid .docx file
            print("\nError: The selected file is not a valid .docx document.")
            print("Tip: Try re-saving the document in Word, then retry.")
            log_error("Invalid docx file", e, file_path)
            return
        except PermissionError as e: # file access permission error
            print("\nError: Permission denied to access the file.")
            print("Tip: Close the document in Word and check file permissions.")
            log_error("Permission error", e, file_path)
            return
        except Exception as e: # other errors
            print("\nError: Unexpected issue while validating the file.")
            print("Tip: Try again or choose a different .docx file.")
            log_error("File validation error", e, file_path)
        return

def read_document(file_path):
    """Read the .docx file and return it as an object"""
    try:
        selected_document = Document(file_path)
    except Exception as e:
        print("\nError: Unable to open the .docx file.")
        print("Tip: Ensure the file is a valid .docx and not corrupted.")
        log_error("Open document error", e, file_path)
        return
    doc = [] # create empty list to store paragraphs

    # extract all text paragraphs from the document
    try:
        print("\nReading document content...")
        for paragraph in selected_document.paragraphs:
#            if paragraph.text.strip() != "": # skip empty paragraphs
            doc.append(paragraph.text)
        # print document content read success message
        print("\nDocument content read successfully.")
    except Exception as e: # handle errors while reading document
        print("\nError: Failed while reading the document content.")
        print("Tip: Try re-saving the document and retry.")
        log_error("Read document error", e, file_path)
        return

    # print paragraph count
    total = len(selected_document.paragraphs)
    print(f"Paragraph count: {total}\n")

    # print all paragraphs with a paragraph index as test
#    for index, paragraph in enumerate(doc):
#        if paragraph.strip() != "": # skip empty paragraphs
#            print(index + 1,
#                  paragraph, f"Selected style: {selected_document.paragraphs[index].style.name}",
#                  f"Alignment: {selected_document.paragraphs[index].alignment}",
#                  f"Font: {selected_document.paragraphs[index].runs[0].font.name if selected_document.paragraphs[index].runs else "Default Font"}",
#                 f"{len(selected_document.paragraphs[index].text)} characters",
#                  sep = " - ")
#            print(f"P{index + 1}: {paragraph}") # alternative print format
#        else:
#            print(index + 1,"<Empty paragraph>", sep=" - ")
#            index - 1 # do not count empty paragraphs
    return selected_document if selected_document else None

def iter_document_blocks(document):
    """Yield body items in document order as ('paragraph'|'table', object)."""
    for child in document.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield "paragraph", Paragraph(child, document)
        elif child.tag.endswith('}tbl'):
            yield "table", Table(child, document)

def iter_container_blocks(container):
    """Yield paragraph/table blocks in order for document, header, or footer containers."""
    if hasattr(container, "element") and hasattr(container.element, "body"):
        parent_elm = container.element.body
    else:
        parent_elm = container._element

    for child in parent_elm.iterchildren():
        if child.tag.endswith('}p'):
            yield "paragraph", Paragraph(child, container)
        elif child.tag.endswith('}tbl'):
            yield "table", Table(child, container)

def count_paragraphs_in_table(table):
    count = 0
    for row in table.rows:
        for cell in row.cells:
            count += len(cell.paragraphs)
    return count

def count_paragraphs_in_container(container):
    count = len(container.paragraphs)
    for table in container.tables:
        count += count_paragraphs_in_table(table)
    return count

def count_paragraphs_in_document(document):
    count = len(document.paragraphs)
    for table in document.tables:
        count += count_paragraphs_in_table(table)

    visited = set()
    for section in document.sections:
        containers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for container in containers:
            key = id(container._element)
            if key in visited:
                continue
            visited.add(key)
            count += count_paragraphs_in_container(container)

    return count

def iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph

def iter_paragraphs_in_container(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for paragraph in iter_paragraphs_in_table(table):
            yield paragraph

def collect_all_paragraphs(document):
    paragraphs = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph)
    for table in document.tables:
        paragraphs.extend(list(iter_paragraphs_in_table(table)))

    visited = set()
    for section in document.sections:
        containers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for container in containers:
            key = id(container._element)
            if key in visited:
                continue
            visited.add(key)
            for paragraph in iter_paragraphs_in_container(container):
                paragraphs.append(paragraph)

    return paragraphs

def count_tables_in_document(document):
    count = len(document.tables)
    visited = set()
    for section in document.sections:
        containers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for container in containers:
            key = id(container._element)
            if key in visited:
                continue
            visited.add(key)
            count += len(container.tables)
    return count

def count_images_in_document(document):
    count = 0
    for node in document.element.body.iter():
        if node.tag.endswith('}drawing') or node.tag.endswith('}pict'):
            count += 1
    return count

def format_duration(seconds):
    seconds = int(seconds)
    minutes = seconds // 60
    seconds = seconds % 60
    return f"{minutes:02d}:{seconds:02d}"

def confirm_success(output_path, stats, open_mode="file", output_dir=None):
    print("\nTranslation completed successfully.")
    print(f"Output file: {output_path}")
    print("Summary:")
    print(f"- Paragraphs translated: {stats['paragraphs']}")
    print(f"- Tables processed: {stats['tables']}")
    print(f"- Images preserved: {stats['images']}")
    print(f"- Headers/Footers processed: {stats['headers_footers']}")
    print(f"- Total time: {stats['elapsed']}")

    try:
        if open_mode == "file":
            answer = input("\nOpen translated file now? (y/n): ").strip().lower()
            if answer == "y":
                print(f"Opening file: {output_path}")
                webbrowser.open(output_path)
        elif open_mode == "folder" and output_dir:
            answer = input("\nOpen translated files folder now? (y/n): ").strip().lower()
            if answer == "y":
                print(f"Opening folder: {output_dir}")
                webbrowser.open(output_dir)
    except Exception:
        pass

def run_validation_checks(original_file_path, output_path, original_paragraphs, original_text_flags):
    warnings = []

    # Check paragraph counts
    current_paragraphs = collect_all_paragraphs(Document(output_path))
    if len(current_paragraphs) != len(original_paragraphs):
        warnings.append(
            f"Paragraph count changed (original {len(original_paragraphs)} vs output {len(current_paragraphs)})."
        )

    # Check for empty translated paragraphs where source had text
    for idx, (para, had_text) in enumerate(zip(current_paragraphs, original_text_flags), start=1):
        if had_text and not para.text.strip():
            warnings.append(f"Paragraph #{idx} became empty after translation.")

    # Compare file sizes
    try:
        src_size = os.path.getsize(original_file_path)
        out_size = os.path.getsize(output_path)
        if src_size > 0:
            diff_ratio = abs(out_size - src_size) / src_size
            if diff_ratio > 0.35:
                warnings.append(
                    f"File size difference is large ({src_size} bytes vs {out_size} bytes)."
                )
    except Exception as e:
        log_error("File size comparison error", e, original_file_path)

    if warnings:
        print("\nValidation warnings:")
        for w in warnings:
            print(f"- {w}")
        log_error("Validation warnings", "\n".join(warnings), original_file_path)
    else:
        print("\nValidation checks passed.")

def run_quality_checks(original_texts, output_paragraphs, log_to_file=True, ui=None):
    warnings = []

    for idx, (src_text, out_para) in enumerate(zip(original_texts, output_paragraphs), start=1):
        src = (src_text or "").strip()
        out = (out_para.text or "").strip()

        if not src:
            continue

        # Flag suspicious length ratios.
        if len(src) > 0:
            ratio = len(out) / max(len(src), 1)
            if ratio < 0.5 or ratio > 2.0:
                msg = (
                    ui["warn_length_ratio"].format(idx=idx, src=len(src), out=len(out))
                    if ui else
                    f"Paragraph #{idx} length ratio suspicious (src {len(src)} vs out {len(out)})."
                )
                warnings.append(msg)

        # Flag common English words that remain.
        out_words = re.findall(r"[A-Za-z']+", out.lower())
        if any(w in COMMON_ENGLISH_WORDS for w in out_words):
            msg = (
                ui["warn_english_words"].format(idx=idx)
                if ui else
                f"Paragraph #{idx} contains common English words."
            )
            warnings.append(msg)

        # Known translation checks (simple substring match).
        src_l = src.lower()
        out_l = out.lower()
        for k, v in KNOWN_TRANSLATIONS.items():
            if k in src_l and v not in out_l:
                msg = (
                    ui["warn_known_translation"].format(idx=idx, phrase=k)
                    if ui else
                    f"Paragraph #{idx} may have missed known translation for '{k}'."
                )
                warnings.append(msg)

    if warnings:
        print("\nQuality warnings:")
        for w in warnings:
            print(f"- {w}")
        if log_to_file:
            log_error("Quality warnings", "\n".join(warnings))
    else:
        print("\nQuality checks passed.")

    return warnings

class ProgressTracker:
    def __init__(self, total):
        self.total = max(total, 1)
        self.completed = 0
        self.start_ts = time.monotonic()

    def tick(self):
        self.completed += 1
        elapsed = time.monotonic() - self.start_ts
        rate = self.completed / elapsed if elapsed > 0 else 0
        remaining = self.total - self.completed
        eta = remaining / rate if rate > 0 else 0
        eta_min = int(eta // 60)
        eta_sec = int(eta % 60)
        pct = (self.completed / self.total) * 100
        print(
            f"Processing... {self.completed}/{self.total} "
            f"({pct:.1f}%) ETA {eta_min:02d}:{eta_sec:02d}",
            end="\r",
            flush=True,
        )

    def finish(self):
        print()

    def elapsed(self):
        return time.monotonic() - self.start_ts

class BatchProgressTracker:
    def __init__(self, total_files):
        self.total = max(total_files, 1)
        self.completed = 0
        self.start_ts = time.monotonic()

    def tick(self):
        self.completed += 1
        elapsed = time.monotonic() - self.start_ts
        rate = self.completed / elapsed if elapsed > 0 else 0
        remaining = self.total - self.completed
        eta = remaining / rate if rate > 0 else 0
        eta_min = int(eta // 60)
        eta_sec = int(eta % 60)
        pct = (self.completed / self.total) * 100
        print(
            f"Batch progress: {self.completed}/{self.total} "
            f"({pct:.1f}%) ETA {eta_min:02d}:{eta_sec:02d}",
            end="\r",
            flush=True,
        )

    def finish(self):
        print()
async def translate_text_preserving_whitespace(text, target_lang):
    """Translate text while preserving whitespace-only chunks unchanged."""
    if text is None or text == "":
        return text
    if text.strip() == "":
        return text

    leading_match = re.match(r"^\s*", text)
    trailing_match = re.search(r"\s*$", text)
    leading_ws = leading_match.group(0) if leading_match else ""
    trailing_ws = trailing_match.group(0) if trailing_match else ""
    core_end = len(text) - len(trailing_ws)
    core_text = text[len(leading_ws):core_end]

    if core_text == "":
        return text

    protected_text, replacements = protect_non_translatables(core_text)
    result_text = await provider_translate(protected_text, target_lang)
    await asyncio.sleep(delay_between_requests)
    translated_core = restore_non_translatables(result_text, replacements)
    translated_core = normalize_unicode(translated_core)
    return f"{leading_ws}{translated_core}{trailing_ws}"

def protect_non_translatables(text):
    """
    Replace non-translatable fragments with placeholders and return:
    (protected_text, replacements_dict)
    """
    protected = text
    replacements = {}
    token_index = 0

    for pattern in NON_TRANSLATABLE_PATTERNS:
        def _replace(match):
            nonlocal token_index
            token = f"__NTX_{token_index}__"
            token_index += 1
            replacements[token] = match.group(0)
            return token

        protected = pattern.sub(_replace, protected)

    # Protect ALL CAPS acronyms anywhere.
    def _replace_caps(match):
        nonlocal token_index
        token = f"__NTX_{token_index}__"
        token_index += 1
        replacements[token] = match.group(0)
        return token

    protected = ALL_CAPS_RE.sub(_replace_caps, protected)

    # Protect explicit words from list (e.g., names that get mistranslated).
    if PROTECTED_WORDS:
        # Longer phrases first to avoid partial matches.
        for phrase in sorted(PROTECTED_WORDS, key=len, reverse=True):
            if not phrase.strip():
                continue
            # Allow flexible whitespace inside multi-word phrases.
            parts = [re.escape(p) for p in phrase.split()]
            pattern = r"\b" + r"\s+".join(parts) + r"\b"
            protected_words_re = re.compile(pattern, re.IGNORECASE)

            def _replace_protected(match):
                nonlocal token_index
                if PLACEHOLDER_RE.search(match.group(0)):
                    return match.group(0)
                token = f"__NTX_{token_index}__"
                token_index += 1
                replacements[token] = match.group(0)
                return token

            protected = protected_words_re.sub(_replace_protected, protected)

    # Protect proper nouns (capitalized words), with sentence-start exception.
    # Multi-word capitalized sequences are always protected (e.g., "John Smith").
    for match in list(PROPER_NOUN_SEQUENCE_RE.finditer(protected)):
        span_text = match.group(0)
        if PLACEHOLDER_RE.search(span_text):
            continue

        words = span_text.split()
        is_multi_word = len(words) > 1

        if not is_multi_word:
            # Check if this single capitalized word is at sentence start.
            idx = match.start()
            j = idx - 1
            while j >= 0 and protected[j].isspace():
                j -= 1
            if j < 0:
                continue
            if protected[j] in ".!?":
                continue

        token = f"__NTX_{token_index}__"
        token_index += 1
        replacements[token] = span_text
        protected = protected[:match.start()] + token + protected[match.end():]

    return protected, replacements

def restore_non_translatables(text, replacements):
    """Restore protected placeholders to their original values."""
    restored = text
    for token, original in replacements.items():
        restored = restored.replace(token, original)
    return restored

def normalize_unicode(text):
    """Normalize text to NFC to keep composed Spanish characters (ñ, á, é, í, ó, ú)."""
    return unicodedata.normalize("NFC", text)

def fix_run_boundary_punctuation_from_texts(source_texts, target_runs):
    """
    Remove punctuation introduced at run boundaries when it did not exist in source.
    This avoids artifacts like: "ella. objetivos" when source boundary had no dot.
    """
    limit = min(len(source_texts), len(target_runs))

    for i in range(1, limit):
        src_prev = source_texts[i - 1].rstrip()
        src_next = source_texts[i].lstrip()
        tgt_prev = target_runs[i - 1]

        if not src_prev or not src_next:
            continue

        src_prev_ended_with_dot = src_prev.endswith(".")
        src_next_starts_word = src_next[0].isalpha() or src_next[0].isdigit()

        tgt_prev_text = tgt_prev.text
        tgt_prev_stripped = tgt_prev_text.rstrip()
        tgt_prev_has_dot = tgt_prev_stripped.endswith(".")

        if (not src_prev_ended_with_dot) and src_next_starts_word and tgt_prev_has_dot:
            # Remove only one trailing period and keep any trailing whitespace.
            m = re.search(r"\.(\s*)$", tgt_prev_text)
            if m:
                tgt_prev.text = tgt_prev_text[:m.start()] + m.group(1)

def is_page_number_run(run):
    """
    Detect runs that belong to page-number fields (PAGE/NUMPAGES/SECTIONPAGES).
    These runs must not be translated.
    """
    page_field_markers = ("PAGE", "NUMPAGES", "SECTIONPAGES")
    word_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # 1) <w:fldSimple w:instr="...PAGE...">
    parent = run._r.getparent()
    while parent is not None:
        if parent.tag.endswith('}fldSimple'):
            instr = (parent.get(f"{word_ns}instr") or "").upper()
            if any(marker in instr for marker in page_field_markers):
                return True
        parent = parent.getparent()

    # 2) Complex field instructions inside run descendants (<w:instrText>)
    for node in run._r.iter():
        if node.tag.endswith('}instrText'):
            instr_text = (node.text or "").upper()
            if any(marker in instr_text for marker in page_field_markers):
                return True

    return False

def has_field_code_nodes(run):
    """True when run contains field code XML nodes that must not be modified."""
    for node in run._r.iter():
        if node.tag.endswith('}fldChar') or node.tag.endswith('}instrText'):
            return True
    return False

def is_image_run(run):
    """Detect runs that contain drawings/pictures to avoid altering image anchors."""
    for node in run._r.iter():
        if node.tag.endswith('}drawing') or node.tag.endswith('}pict'):
            return True
    return False

async def translate_paragraph_in_place(paragraph, target_lang, progress=None):
    """Translate one paragraph in place, preserving list/paragraph/run formatting."""
    if paragraph.runs:
        source_run_texts = [run.text for run in paragraph.runs]
        in_field = False

        for idx, run in enumerate(paragraph.runs):
            if is_image_run(run):
                continue
            # Never rewrite field-code runs; it can break PAGE/NUMPAGES rendering.
            if has_field_code_nodes(run):
                for node in run._r.iter():
                    if node.tag.endswith('}fldChar'):
                        fld_char_type = node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType")
                        if fld_char_type == "begin":
                            in_field = True
                        elif fld_char_type == "end":
                            in_field = False
                continue

            # Keep all field-result runs untouched while inside a field.
            if in_field or is_page_number_run(run):
                continue
            source_text = source_run_texts[idx]
            run.text = await translate_text_preserving_whitespace(source_text, target_lang)

        fix_run_boundary_punctuation_from_texts(source_run_texts, paragraph.runs)
        if progress:
            progress.tick()
        return

    if paragraph.text and paragraph.text.strip():
        paragraph.text = await translate_text_preserving_whitespace(paragraph.text, target_lang)
    if progress:
        progress.tick()

async def translate_table_in_place(table, target_lang, progress=None):
    """Translate text in each table cell paragraph in place."""
    print(f"Translating table: {len(table.rows)} row(s), {len(table.columns)} col(s)")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                await translate_paragraph_in_place(paragraph, target_lang, progress)

async def translate_header_footer_in_place(document, target_lang, progress=None):
    """Translate all unique headers and footers in place, preserving page-number fields."""
    visited = set()

    for section in document.sections:
        containers = (
            ("header_default", section.header),
            ("header_first", section.first_page_header),
            ("header_even", section.even_page_header),
            ("footer_default", section.footer),
            ("footer_first", section.first_page_footer),
            ("footer_even", section.even_page_footer),
        )

        for container_type, container in containers:
            key = (container_type, id(container._element))
            if key in visited:
                continue
            visited.add(key)

            for block_type, block in iter_container_blocks(container):
                if block_type == "paragraph":
                    await translate_paragraph_in_place(block, target_lang, progress)
                elif block_type == "table":
                    await translate_table_in_place(block, target_lang, progress)

async def translated_doc_creation(file_path, selected_document, target_lang="ES", open_mode="file", ui_strings=None):
    """
    Function to create an output file and write translated text into the new file.
    Returns:
    - output file path if successful or None if failed
    - prints error messages if any error occurs
    - creates a new .docx file with translated text and preserved formatting
    """
    print("\nCreating output file...")

    # Translate in place on source structure to preserve numbering/list definitions.
    trans_file = selected_document

    # Select original file path
    try:
        base_name = os.path.basename(file_path)
        name_no_ext, ext = os.path.splitext(base_name)
        print()
        print(f"Chosen base name: {base_name}")
    except Exception as e:
        print("\nError: Could not read the selected file name.")
        print("Tip: Check the file name and try again.")
        log_error("Read file name error", e, file_path)
        return None, None, None, []

    try:
        output_dir = (os.getenv("lin_translated_docs_dir")
        or os.getenv("translated_docs_dir")
        or os.path.dirname(file_path))

        # check if output directory is defined
        if not output_dir:
            print("\nError: No output directory defined.")
            print("Tip: Set 'translated_docs_dir' or choose a file in a writable folder.")
            log_error("Output directory missing", "No output directory", file_path)
            return None, None, None, []

        safe_base = "".join(name_no_ext.split())
        safe_base = re.sub(r'[\\/:\*\?"<>\|]', "", safe_base)
        output_name = f"{safe_base}_{target_lang.upper()}{ext}"

        output_path = os.path.join(output_dir,
                                output_name)
        print(f"Chosen file dir: {output_dir}")
        print(f"New file name: {output_name}")
        print(f"New file path: {output_path}\n")
    except Exception as e:
        print("\nError: Could not determine output directory.")
        print("Tip: Check environment variables or file permissions.")
        log_error("Output directory error", e, file_path)
        return None, None, None, []

    try:
        total_paragraphs = count_paragraphs_in_document(trans_file)
        total_tables = count_tables_in_document(trans_file)
        total_images = count_images_in_document(trans_file)
        headers_footers = len(trans_file.sections) * 2
        print(f"Processing... Total paragraphs: {total_paragraphs}")
        progress = ProgressTracker(total_paragraphs)

        original_paragraphs = collect_all_paragraphs(trans_file)
        original_text_flags = [bool(p.text.strip()) for p in original_paragraphs]
        original_texts = [p.text for p in original_paragraphs]

        for block_type, block in iter_document_blocks(trans_file):
            if block_type == "paragraph":
                await translate_paragraph_in_place(block, target_lang, progress)
            elif block_type == "table":
                await translate_table_in_place(block, target_lang, progress)

        await translate_header_footer_in_place(trans_file, target_lang, progress)
        progress.finish()

        trans_file.save(output_path)
        save_translation_memory()
        stats = {
            "paragraphs": total_paragraphs,
            "tables": total_tables,
            "images": total_images,
            "headers_footers": headers_footers,
            "elapsed": format_duration(progress.elapsed()),
        }
        run_validation_checks(file_path, output_path, original_paragraphs, original_text_flags)
        quality_warnings = []
        try:
            output_paragraphs = collect_all_paragraphs(Document(output_path))
            quality_warnings = run_quality_checks(
                original_texts,
                output_paragraphs,
                log_to_file=(open_mode != "none"),
                ui=ui_strings
            )
        except Exception as e:
            log_error("Quality check error", e, file_path)
        confirm_success(output_path, stats, open_mode=open_mode, output_dir=output_dir)
        return output_path, output_dir, stats, quality_warnings

    except Exception as e:
            print("\nError: Translation failed while writing the output file.")
            print("Tip: Check disk space and file permissions, then try again.")
            log_error("Translation/output error", e, file_path)
            return None, None, None, []

def inspect_tables(selected_document):
    """"""
    if not selected_document.tables:
        print(f'\nNo tables found in the document.')
        return

    print(f'\nFound {len(selected_document.tables)} tables.\n')

 #   for t_idx, table in enumerate(selected_document.tables):
 #       t_idx += 1
 #       print(f'Table {t_idx}:',
 #             f'    Rows: {len(table.rows)}')

 #       for r_idx, row in enumerate(table.rows):
 #           print(f'        Row {r_idx + 1}: Cells = {len(row.cells)}')

 #           for c_idx, cell in enumerate(row.cells):
 #               print(f'            Cell {c_idx + 1}: Paragraphs = {len(cell.paragraphs)}')

  #              for p_idx, paragraph in enumerate(cell.paragraphs):
  #                  text = paragraph.text.strip()
  #                  print(f'                Paragraph {p_idx + 1}: "{text}"')

  #      print('-'*20, f'Table {t_idx} end', '-'*20) # spacing between tables

# the following instances are for debugging and learning purposes only
def debug_print_runs(selected_document):
    """
    Debug function to print all runs in a paragraph.
    Returns:
    - prints all runs in each paragraph with their formatting details
    """
    print()
    print("-"*20, "Run Debug Start", "-"*20)

    for p_idx, paragraph in enumerate(selected_document.paragraphs, start=1):
        print(f"Paragraph {p_idx}:")
        print(f"Full text: {repr(paragraph.text)}")
        print(f"Run count: {len(paragraph.runs)}")

        for r_idx, run in enumerate(paragraph.runs, start=1):
            print(f"    Run {r_idx}: {repr(run.text)} | "
                  f"bold={run.bold}, italic={run.italic}, "
                  f"underline={run.underline}, "
                  f"font={run.font.name}, size={run.font.size}, "
                  )

        print("-"*50)
    print("-"*20, "Run Debug End", "-"*20)

# the following instance is for debugging and learning purposes only
def debug_paragraph_runs(run_info):
    """
    Debug function to print runs of all paragraph
    Returns:
    - prints all runs in each paragraph with their formatting details
    """
    print()
    print("-"*20, "Paragraphs Run Info Start", "-"*20)

    if not run_info:
        print("No paragraphs found.")
        return

    for p_idx, paragraph in enumerate(run_info, start=1):
        print(f"Paragraph {p_idx} -")
        print(f"Alignment: {paragraph['alignment']}")
        print(f"Style: {paragraph['style']}")
        print(f"Run count: {len(paragraph['runs'])}")

        for r_idx, run in enumerate(paragraph["runs"], start=1):
            print(f"Run {r_idx}:")
            for key, value in run.items():
                print(f"    {key}: {value}")
            print()
#        break  # only the first paragraph
    print("-"*20, "All Paragraphs Run Info End", "-"*20)

async def main():
    settings, settings_path = load_settings()
    init_translation_memory(settings, Path(settings_path).parent)
    if settings.get("open_settings_on_start"):
        prompt_edit_settings(settings_path)

    print("Select .docx files to translate...")

    selected_files = []

    while True:
        chosen_file = select_docx_file()
        if not chosen_file:
            if not selected_files:
                print("\nUser closed the window before selecting a file.",
                      "\nGoodbye!")
                return
            break

        selected_files.append(chosen_file)
        if not prompt_add_another_file():
            break

    total_files = len(selected_files)
    print(f"\nBatch start: {total_files} file(s) queued.")

    successes = 0
    failures = 0
    batch_errors = []
    batch_progress = BatchProgressTracker(total_files)
    batch_output_dir = None

    for idx, file_path in enumerate(selected_files, start=1):
        print(f"\nProcessing file {idx}/{total_files}: {file_path}")

        # validate file
        if file_validation(file_path) is None:
            failures += 1
            batch_errors.append((file_path, "Validation failed"))
            continue

        # read document
        selected_document = read_document(file_path)
        if not selected_document:
            failures += 1
            batch_errors.append((file_path, "Read failed"))
            continue

        # inspect for tables
        inspect_tables(selected_document)

        try:
            open_mode = "file" if total_files == 1 else "none"
            target_lang = validate_language_code(settings.get("target_lang", "ES"), "ES")
            output_path, output_dir, _stats, _quality_warnings = await translated_doc_creation(file_path,
                                                                                               selected_document,
                                                                                               target_lang=target_lang,
                                                                                               open_mode=open_mode)
            if output_path:
                successes += 1
                if total_files > 1 and output_dir:
                    batch_output_dir = output_dir
            else:
                failures += 1
                batch_errors.append((file_path, "Translation failed"))
        except Exception as e:
            failures += 1
            batch_errors.append((file_path, f"Exception: {e}"))
            log_error("Batch translation error", e, file_path)
            continue
        finally:
            batch_progress.tick()

    batch_progress.finish()
    print("\nBatch completed.")
    print(f"- Successes: {successes}")
    print(f"- Failures: {failures}")
    if batch_errors:
        print("Failed files:")
        for fp, reason in batch_errors:
            print(f"- {fp} | {reason}")
    if total_files > 1 and successes > 0 and batch_output_dir:
        try:
            answer = input("\nOpen translated files folder now? (y/n): ").strip().lower()
            if answer == "y":
                print(f"Opening folder: {batch_output_dir}")
                webbrowser.open(batch_output_dir)
        except Exception:
            pass
    #if output_path:
    #    os.system(f"open {output_path}")

def gui_main():
    settings, settings_path = load_settings()
    init_translation_memory(settings, Path(settings_path).parent)
    show_language_selector(settings, settings_path)
    ui = get_ui_strings(settings)

    root = Tk()
    root.title(ui["app_title"])
    root.geometry("720x520")

    selected_files = []
    stop_flag = {"stop": False}
    ui_queue = queue.Queue()

    def add_files():
        files = fd.askopenfilenames(
            title="Choose .docx files",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if not files:
            return
        for f in files:
            if f not in selected_files:
                selected_files.append(f)
                listbox.insert("end", f)

    def clear_files():
        selected_files.clear()
        listbox.delete(0, "end")
        progress_bar["value"] = 0
        status_var.set(ui["ready"])
        summary_text.delete("1.0", "end")
        start_btn.config(state="normal")
        stop_btn.config(state="disabled")
        stop_flag["stop"] = False

    def on_open_settings():
        win = tk.Toplevel(root)
        win.title(ui["settings_title"])
        win.geometry("420x260")

        ttk.Label(win, text=ui["target_language"]).pack(anchor="w", padx=10, pady=5)

        languages = [
            ("Spanish", "ES"),
            ("English", "EN"),
            ("Portuguese", "PT"),
            ("Italian", "IT"),
            ("French", "FR"),
            ("Chinese", "ZH-CN"),
        ]

        lang_var = tk.StringVar(value=settings.get("target_lang", "ES"))
        lang_menu = ttk.Combobox(
            win,
            values=[f"{label} ({code})" for label, code in languages],
            state="readonly"
        )
        code_to_label = {code: f"{label} ({code})" for label, code in languages}
        if lang_var.get() in code_to_label:
            lang_menu.set(code_to_label[lang_var.get()])
        else:
            lang_menu.set("Spanish (ES)")
        lang_menu.pack(fill="x", padx=10)

        open_settings_var = tk.BooleanVar(value=bool(settings.get("open_settings_on_start")))
        ttk.Checkbutton(
            win,
            text=ui["open_settings_start"],
            variable=open_settings_var
        ).pack(anchor="w", padx=10, pady=10)

        ttk.Label(win, text=ui["language_prompt"]).pack(anchor="w", padx=10, pady=5)
        ui_lang_var = tk.StringVar(value=settings.get("ui_language", "EN"))
        ui_lang_menu = ttk.Combobox(
            win,
            values=[f"{ui['english']} (EN)", f"{ui['spanish']} (ES)"],
            state="readonly"
        )
        if ui_lang_var.get().upper() == "ES":
            ui_lang_menu.set(f"{ui['spanish']} (ES)")
        else:
            ui_lang_menu.set(f"{ui['english']} (EN)")
        ui_lang_menu.pack(fill="x", padx=10)

        use_gui_var = tk.BooleanVar(value=bool(settings.get("use_gui")))
        ttk.Checkbutton(
            win,
            text=ui["use_gui"],
            variable=use_gui_var
        ).pack(anchor="w", padx=10, pady=5)

        def save_and_close():
            previous_ui_lang = settings.get("ui_language", "EN")
            selection = lang_menu.get()
            for label, code in languages:
                if selection.startswith(label):
                    settings["target_lang"] = code
                    break
            settings["open_settings_on_start"] = open_settings_var.get()
            settings["use_gui"] = use_gui_var.get()
            ui_selection = ui_lang_menu.get()
            if ui_selection.endswith("(ES)"):
                settings["ui_language"] = "ES"
            else:
                settings["ui_language"] = "EN"
            save_settings(settings, settings_path)
            if settings["ui_language"] != previous_ui_lang:
                root.destroy()
                gui_main()
            win.destroy()

        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill="x", padx=10, pady=10)
        ttk.Button(btn_frame, text=ui["save"], command=save_and_close).pack(side="right")
        ttk.Button(btn_frame, text=ui["cancel"], command=win.destroy).pack(side="right", padx=5)

    def set_status(msg):
        status_var.set(msg)

    def worker():
        if not selected_files:
            ui_queue.put(("status", "No files selected."))
            return

        stop_flag["stop"] = False
        successes = 0
        failures = 0
        batch_errors = []
        batch_output_dir = None
        total_files = len(selected_files)

        ui_queue.put(("batch_total", total_files))
        ui_queue.put(("status", ui["batch_start"].format(count=total_files)))

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        local_translator = Translator()
        set_translator(local_translator)

        for idx, file_path in enumerate(selected_files, start=1):
            if stop_flag["stop"]:
                ui_queue.put(("status", "Batch stopped by user."))
                break

            ui_queue.put(("status", ui["processing_file"].format(idx=idx, total=total_files, file=file_path)))
            ui_queue.put(("file_start", (idx, total_files, file_path)))

            if file_validation(file_path) is None:
                failures += 1
                batch_errors.append((file_path, "Validation failed"))
                ui_queue.put(("batch_tick", None))
                continue

            selected_document = read_document(file_path)
            if not selected_document:
                failures += 1
                batch_errors.append((file_path, "Read failed"))
                ui_queue.put(("batch_tick", None))
                continue

            inspect_tables(selected_document)

            try:
                target_lang = validate_language_code(settings.get("target_lang", "ES"), "ES")
                output_path, output_dir, stats, quality_warnings = loop.run_until_complete(
                    translated_doc_creation(
                        file_path,
                        selected_document,
                        target_lang=target_lang,
                        open_mode="none",
                        ui_strings=ui
                    )
                )
                if output_path:
                    successes += 1
                    batch_output_dir = output_dir
                    ui_queue.put(("summary", (file_path, output_path, stats, quality_warnings)))
                    ui_queue.put(("file_done", (idx, total_files)))
                else:
                    failures += 1
                    batch_errors.append((file_path, "Translation failed"))
                    ui_queue.put(("file_done", (idx, total_files)))
            except Exception as e:
                failures += 1
                batch_errors.append((file_path, f"Exception: {e}"))
                log_error("GUI batch translation error", e, file_path)
                ui_queue.put(("file_done", (idx, total_files)))
            finally:
                ui_queue.put(("batch_tick", None))

        save_translation_memory()
        loop.run_until_complete(close_translator_client(local_translator))
        loop.close()
        ui_queue.put(("batch_done", (successes, failures, batch_errors, batch_output_dir)))

    def start_batch():
        if not selected_files:
            messagebox.showwarning("No files", ui["no_files"])
            return
        output_dir = select_output_folder()
        if not output_dir:
            return
        settings["translated_docs_dir"] = output_dir
        os.environ["translated_docs_dir"] = output_dir
        start_btn.config(state="disabled")
        stop_btn.config(state="normal")
        thread = threading.Thread(target=worker, daemon=True)
        thread.start()

    def stop_batch():
        stop_flag["stop"] = True
        stop_btn.config(state="disabled")

    def process_queue():
        try:
            while True:
                item = ui_queue.get_nowait()
                kind = item[0]
                payload = item[1] if len(item) > 1 else None

                if kind == "status":
                    set_status(payload)
                elif kind == "batch_total":
                    progress_bar["maximum"] = payload
                    progress_bar["value"] = 0
                    file_progress["value"] = 0
                elif kind == "file_start":
                    if payload is None:
                        continue
                    idx, total, file_path = payload
                    file_progress["value"] = 0
                    file_status_var.set(f"File {idx}/{total}: {os.path.basename(file_path)}")
                elif kind == "batch_tick":
                    progress_bar["value"] = progress_bar["value"] + 1
                elif kind == "file_done":
                    if payload is None:
                        continue
                    idx, total = payload
                    file_progress["value"] = 100
                    file_status_var.set(f"File {idx}/{total}: done")
                elif kind == "summary":
                    if payload is None:
                        continue
                    file_path, output_path, stats, quality_warnings = payload
                    summary_text.insert("end", f"\n{file_path}\n")
                    summary_text.insert("end", f"{ui['summary_output']}: {output_path}\n")
                    summary_text.insert(
                        "end",
                        f"{ui['summary_paragraphs']}: {stats['paragraphs']} | {ui['summary_tables']}: {stats['tables']} | "
                        f"{ui['summary_images']}: {stats['images']} | {ui['summary_headers']}: {stats['headers_footers']} | "
                        f"{ui['summary_time']}: {stats['elapsed']}\n"
                    )
                    if quality_warnings:
                        summary_text.insert("end", f"{ui['quality_warnings']}:\n")
                        for w in quality_warnings:
                            summary_text.insert("end", f"- {w}\n")
                    summary_text.see("end")
                elif kind == "batch_done":
                    if payload is None:
                        continue
                    successes, failures, batch_errors, output_dir = payload
                    set_status("Batch completed.")
                    start_btn.config(state="normal")
                    stop_btn.config(state="disabled")

                    if batch_errors:
                        msg = "\n".join([f"{fp} | {reason}" for fp, reason in batch_errors])
                        messagebox.showwarning(ui["batch_completed"], msg)
                    else:
                        messagebox.showinfo(ui["batch_completed"], ui["all_files_ok"])

                    if output_dir:
                        if messagebox.askyesno(ui["batch_completed"], ui["open_folder"]):
                            webbrowser.open(output_dir)
                ui_queue.task_done()
        except queue.Empty:
            pass
        root.after(200, process_queue)

    header = ttk.Label(root, text=ui["app_title"], font=("Segoe UI", 16))
    header.pack(pady=10)

    btn_frame = ttk.Frame(root)
    btn_frame.pack(fill="x", padx=10)

    add_btn = ttk.Button(btn_frame, text=ui["add_files"], command=add_files)
    add_btn.pack(side="left", padx=5)

    clear_btn = ttk.Button(btn_frame, text=ui["clear_list"], command=clear_files)
    clear_btn.pack(side="left", padx=5)

    settings_btn = ttk.Button(btn_frame, text=ui["settings"], command=on_open_settings)
    settings_btn.pack(side="right", padx=5)

    listbox = tk.Listbox(root, height=8)
    listbox.pack(fill="both", expand=False, padx=10, pady=10)

    progress_bar = ttk.Progressbar(root, mode="determinate")
    progress_bar.pack(fill="x", padx=10, pady=(0, 5))

    file_progress = ttk.Progressbar(root, mode="determinate", maximum=100)
    file_progress.pack(fill="x", padx=10)

    file_status_var = tk.StringVar(value=ui["file_idle"])
    file_status_label = ttk.Label(root, textvariable=file_status_var)
    file_status_label.pack(fill="x", padx=10, pady=(0, 5))

    status_var = tk.StringVar(value="Ready.")
    status_label = ttk.Label(root, textvariable=status_var)
    status_label.pack(fill="x", padx=10, pady=5)

    control_frame = ttk.Frame(root)
    control_frame.pack(fill="x", padx=10, pady=5)

    start_btn = ttk.Button(control_frame, text=ui["start"], command=start_batch)
    start_btn.pack(side="left", padx=5)

    stop_btn = ttk.Button(control_frame, text=ui["stop"], command=stop_batch, state="disabled")
    stop_btn.pack(side="left", padx=5)

    summary_label = ttk.Label(root, text=ui["summary"])
    summary_label.pack(anchor="w", padx=10)

    summary_frame = ttk.Frame(root)
    summary_frame.pack(fill="both", expand=True, padx=10, pady=5)

    summary_scroll = ttk.Scrollbar(summary_frame, orient="vertical")
    summary_scroll.pack(side="right", fill="y")

    summary_text = tk.Text(root, height=8, wrap="word", yscrollcommand=summary_scroll.set)
    summary_text.pack(in_=summary_frame, fill="both", expand=True)
    summary_scroll.config(command=summary_text.yview)

    process_queue()
    root.mainloop()

if __name__=="__main__":
    settings, _settings_path = load_settings()
    if settings.get("use_gui"):
        gui_main()
    else:
        asyncio.run(main())
